"""Extract text content from various file formats."""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Max characters to extract per file
MAX_CHARS = 50_000


def extract_content(filename: str, data: bytes) -> tuple[str, str]:
    """
    Extract text content from file data.

    Returns: (content, status_msg)
    - content: extracted text (truncated if > MAX_CHARS)
    - status_msg: brief info about extraction (e.g., "50KB PDF, 25K chars")
    """
    ext = Path(filename).suffix.lower()

    try:
        if ext == ".pdf":
            return _extract_pdf(data, filename)
        elif ext == ".docx":
            return _extract_docx(data, filename)
        elif ext in (".xlsx", ".xls"):
            return _extract_xlsx(data, filename)
        else:
            # Try UTF-8 decode for text files (md, txt, py, js, json, yaml, csv, etc)
            return _extract_text(data, filename)
    except Exception as e:
        logger.error(f"Error extracting {filename}: {e}")
        return "", f"Error reading file: {e}"


def _extract_pdf(data: bytes, filename: str) -> tuple[str, str]:
    """Extract text from PDF."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "pypdf not installed. Run: pip install pypdf"

    try:
        pdf = PdfReader(io.BytesIO(data))
        text = ""
        for page in pdf.pages:
            text += page.extract_text() or ""

        text = text[:MAX_CHARS]
        was_truncated = len(pdf.pages) > 0 and len(text) == MAX_CHARS
        msg = f"PDF ({len(pdf.pages)} pages), {len(text)} chars"
        if was_truncated:
            msg += " [truncated]"
        return text, msg
    except Exception as e:
        return "", f"PDF error: {e}"


def _extract_docx(data: bytes, filename: str) -> tuple[str, str]:
    """Extract text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        return "", "python-docx not installed. Run: pip install python-docx"

    try:
        doc = Document(io.BytesIO(data))
        text = ""

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                text += row_text + "\n"

        was_truncated = len(text) > MAX_CHARS
        text = text[:MAX_CHARS]

        msg = f"DOCX ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables), {len(text)} chars"
        if was_truncated:
            msg += " [truncated]"
        return text, msg
    except Exception as e:
        return "", f"DOCX error: {e}"


def _extract_xlsx(data: bytes, filename: str) -> tuple[str, str]:
    """Extract text from XLSX/XLS (convert to CSV-like format)."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", "openpyxl not installed. Run: pip install openpyxl"

    try:
        wb = load_workbook(io.BytesIO(data), data_only=True)
        text = ""

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"=== Sheet: {sheet_name} ===\n"

            for row in ws.iter_rows(values_only=True):
                # Convert row to pipe-separated values
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                text += row_text + "\n"

            text += "\n"

        was_truncated = len(text) > MAX_CHARS
        text = text[:MAX_CHARS]

        msg = f"XLSX ({len(wb.sheetnames)} sheets), {len(text)} chars"
        if was_truncated:
            msg += " [truncated]"
        return text, msg
    except Exception as e:
        return "", f"XLSX error: {e}"


def _extract_text(data: bytes, filename: str) -> tuple[str, str]:
    """Try UTF-8 decode for plain text, code files, etc."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        return "", "File is not UTF-8 text (binary file?)"

    was_truncated = len(text) > MAX_CHARS
    text = text[:MAX_CHARS]

    msg = f"{Path(filename).suffix[1:].upper() or 'text'} ({len(text)} chars)"
    if was_truncated:
        msg += " [truncated]"
    return text, msg
